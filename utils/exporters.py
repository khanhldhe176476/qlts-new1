from io import BytesIO
from typing import Iterable, Dict, List, Optional
from datetime import datetime


def rows_to_dicts(rows: Iterable, fields: List[str]) -> List[Dict[str, str]]:

	result: List[Dict[str, str]] = []
	for r in rows:
		row = {}
		for f in fields:
			row[f] = getattr(r, f, None)
		result.append(row)
	return result


def export_excel(rows: Iterable, fields: List[str], title: str = "Data", header_map: Optional[Dict[str, str]] = None) -> BytesIO:

	# Lazy import heavy deps to avoid slowing app startup
	import pandas as pd  # type: ignore

	buf = BytesIO()
	data = rows_to_dicts(rows, fields)
	df = pd.DataFrame(data, columns=fields)
	# Vietnamese headers
	if header_map:
		vn_headers = [header_map.get(f, f) for f in fields]
		df.columns = vn_headers
	with pd.ExcelWriter(buf, engine='openpyxl') as writer:
		df.to_excel(writer, index=False, sheet_name=title[:31] or 'Sheet1')
	buf.seek(0)
	return buf


def export_docx(rows: Iterable, fields: List[str], title: str = "Data", header_map: Optional[Dict[str, str]] = None, font_name: str = 'Arial') -> BytesIO:

	from docx import Document  # type: ignore
	from docx.oxml.ns import qn  # type: ignore
	from docx.shared import Pt  # type: ignore

	doc = Document()
	# Set default font that supports Vietnamese
	style = doc.styles['Normal']
	style.font.name = font_name
	# For compatibility with some Word versions
	style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
	style.font.size = Pt(11)
	doc.add_heading(title, level=1)
	table = doc.add_table(rows=1, cols=len(fields))
	hdr_cells = table.rows[0].cells
	for idx, f in enumerate(fields):
		hdr_text = header_map.get(f, f) if header_map else f
		hdr_cells[idx].text = hdr_text
	for r in rows:
		row_cells = table.add_row().cells
		for idx, f in enumerate(fields):
			val = getattr(r, f, "")
			text = "" if val is None else str(val)
			row_cells[idx].text = text
			# Ensure font for each cell
			for p in row_cells[idx].paragraphs:
				for run in p.runs:
					run.font.name = font_name
					run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
	buf = BytesIO()
	doc.save(buf)
	buf.seek(0)
	return buf


def export_pdf(rows: Iterable, fields: List[str], title: str = "Data", header_map: Optional[Dict[str, str]] = None,
			   preferred_fonts: Optional[List[str]] = None) -> BytesIO:

	from reportlab.lib import colors  # type: ignore
	from reportlab.lib.pagesizes import A4  # type: ignore
	from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer  # type: ignore
	from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
	from reportlab.pdfbase import pdfmetrics  # type: ignore
	from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
	import os

	# Try to register a Vietnamese-capable system font
	def _register_vn_font() -> str:
		candidates = preferred_fonts or [
			# Common Windows fonts supporting Vietnamese
			'C:/Windows/Fonts/arial.ttf',
			'C:/Windows/Fonts/tahoma.ttf',
			'C:/Windows/Fonts/times.ttf',
			'C:/Windows/Fonts/verdana.ttf',
			# DejaVu (often installed)
			'/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
		]
		for path in candidates:
			try:
				if os.path.exists(path):
					pdfmetrics.registerFont(TTFont('VN_FONT', path))
					return 'VN_FONT'
			except Exception:
				continue
		# Fallback: Helvetica (may not render tone marks correctly)
		return 'Helvetica'

	buf = BytesIO()
	# Build table data
	# Header row in Vietnamese
	headers = [header_map.get(f, f) for f in fields] if header_map else fields
	data = [headers]
	for r in rows:
		data.append(["" if getattr(r, f, None) is None else str(getattr(r, f)) for f in fields])

	doc = SimpleDocTemplate(buf, pagesize=A4)
	styles = getSampleStyleSheet()
	font_name = _register_vn_font()
	# Override default style with Vietnamese-capable font
	styles.add(ParagraphStyle(name='VNHeading', parent=styles['Heading1'], fontName=font_name))
	styles.add(ParagraphStyle(name='VNBody', parent=styles['BodyText'], fontName=font_name))
	story = [Paragraph(title, styles['VNHeading']), Spacer(1, 12)]
	table = Table(data, repeatRows=1)
	table.setStyle(TableStyle([
		('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
		('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
		('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
		('ALIGN', (0, 0), (-1, -1), 'LEFT'),
		('FONTNAME', (0, 0), (-1, 0), font_name),
		('FONTNAME', (0, 1), (-1, -1), font_name),
		('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.beige])
	]))
	story.append(table)
	doc.build(story)
	buf.seek(0)
	return buf


def export_maintenance_to_excel(records: Iterable) -> bytes:
	"""Export maintenance records to Excel"""
	import pandas as pd
	
	data = []
	for record in records:
		asset = record.asset if hasattr(record, 'asset') and record.asset else None
		data.append({
			'Mã tài sản': asset.device_code if asset else '',
			'Tên tài sản': asset.name if asset else '',
			'Loại tài sản': asset.asset_type.name if asset and asset.asset_type else '',
			'Đơn vị sử dụng': asset.user.username if asset and asset.user else (asset.user_text if asset else ''),
			'Trạng thái thiết bị': asset.status if asset else '',
			'Ngày yêu cầu': record.request_date.strftime('%d/%m/%Y') if record.request_date else '',
			'Người yêu cầu': record.requested_by.username if record.requested_by else '',
			'Nguyên nhân': {
				'broken': 'Hỏng hóc',
				'periodic': 'Bảo trì định kỳ',
				'calibration': 'Hiệu chỉnh',
				'other': 'Khác'
			}.get(record.maintenance_reason, record.maintenance_reason or ''),
			'Mô tả tình trạng': record.condition_before or '',
			'Mức độ hỏng': {
				'light': 'Nhẹ',
				'medium': 'Trung bình',
				'heavy': 'Nặng'
			}.get(record.damage_level, record.damage_level or ''),
			'Đơn vị bảo trì': record.vendor or '',
			'Người thực hiện': record.person_in_charge or '',
			'Số điện thoại': record.vendor_phone or '',
			'Chi phí dự kiến': record.estimated_cost or 0,
			'Ngày hoàn thành': record.completed_date.strftime('%d/%m/%Y') if record.completed_date else '',
			'Phụ tùng thay thế': record.replaced_parts or '',
			'Chi phí thực tế': record.cost or 0,
			'Ghi chú': record.result_notes or '',
			'Trạng thái sau bảo trì': {
				'passed': 'Đạt',
				'failed': 'Không đạt',
				'need_retry': 'Cần bảo trì lại'
			}.get(record.result_status, record.result_status or ''),
			'Trạng thái': {
				'pending': 'Chờ xử lý',
				'in_progress': 'Đang thực hiện',
				'completed': 'Hoàn thành',
				'failed': 'Không đạt',
				'cancelled': 'Hủy'
			}.get(record.status, record.status)
		})
	
	df = pd.DataFrame(data)
	buf = BytesIO()
	with pd.ExcelWriter(buf, engine='openpyxl') as writer:
		df.to_excel(writer, sheet_name='Danh sách bảo trì', index=False)
	buf.seek(0)
	return buf.read()

